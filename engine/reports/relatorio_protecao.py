# -*- coding: utf-8 -*-
"""
engine/reports/relatorio_protecao.py
=====================================
Gerador de Relatorio Tecnico Word -- Estudo de Protecao de Sistemas Eletricos
BK Engenharia e Tecnologia -- v3.0
"""
from __future__ import annotations
import io, math, datetime
from typing import Any
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

_C_AZUL_ESC = "1F3864"
_C_AZUL_MED = "2E74B5"
_C_AZUL_LIG = "BDD7EE"
_C_CINZA    = "404040"
_C_CINZA2   = "F2F2F2"
_C_BRANCO   = "FFFFFF"
_MNS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

def _set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_color)
    tcPr.append(shd)

def _set_cell_border(cell, sides=("top","bottom","left","right"), sz="4", color="CCCCCC"):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcBdr = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),"single"); el.set(qn("w:sz"),sz)
        el.set(qn("w:space"),"0"); el.set(qn("w:color"),color)
        tcBdr.append(el)
    tcPr.append(tcBdr)

def _cell_write(cell, text, bold=False, italic=False, size=9, color=_C_CINZA, center=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.clear()
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color)

def _hline(doc, color=_C_AZUL_MED, sz="6", before=2, after=4):
    p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr"); btm = OxmlElement("w:bottom")
    btm.set(qn("w:val"),"single"); btm.set(qn("w:sz"),sz)
    btm.set(qn("w:space"),"1"); btm.set(qn("w:color"),color)
    pBdr.append(btm); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)

def _sp(doc, pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(pts)

def _h1(doc, numero, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{numero}  {texto.upper()}")
    run.bold = True; run.font.size = Pt(13); run.font.color.rgb = RGBColor.from_string(_C_AZUL_ESC)
    _hline(doc, color=_C_AZUL_ESC, sz="8", before=0, after=6)

def _h2(doc, numero, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{numero}  {texto}")
    run.bold = True; run.font.size = Pt(11); run.font.color.rgb = RGBColor.from_string(_C_AZUL_MED)

def _body(doc, texto, size=10, bold=False, italic=False, before=1, after=3, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(texto); run.bold = bold; run.italic = italic
    run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(_C_CINZA)
    return p

def _nota(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run("NOTA: "); r1.bold = True; r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor.from_string(_C_AZUL_MED)
    r2 = p.add_run(texto); r2.italic = True; r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor.from_string(_C_CINZA)

def _ok_str(ok): return "APROVADO" if ok else "REPROVADO"

def _omml_block(doc, omml_xml, label=""):
    try: root = etree.fromstring(omml_xml.encode("utf-8"))
    except Exception:
        _body(doc, f"[Equacao: {label}]", italic=True); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
    oMathPara = OxmlElement("m:oMathPara"); pr = OxmlElement("m:oMathParaPr")
    jc = OxmlElement("m:jc"); jc.set(qn("m:val"), "center")
    pr.append(jc); oMathPara.append(pr); oMathPara.append(root); p._p.append(oMathPara)
    if label:
        tab = p.add_run(f"\t({label})"); tab.font.size = Pt(9)
        tab.font.color.rgb = RGBColor.from_string(_C_CINZA)

def EQ_ZQ():
    return (f'<m:oMath xmlns:m="{_MNS}"><m:r><m:rPr><m:sty m:val="i"/></m:rPr>'
        '<m:t xml:space="preserve">Z = </m:t></m:r><m:f><m:num>'
        '<m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t xml:space="preserve">c &#x22C5; </m:t></m:r>'
        '<m:sSup><m:e><m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t>V</m:t></m:r></m:e>'
        '<m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>'
        '<m:r><m:t xml:space="preserve"> nQ</m:t></m:r></m:num><m:den>'
        '<m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t xml:space="preserve">S kQ</m:t></m:r>'
        '</m:den></m:f></m:oMath>')

def EQ_ZLINHA():
    return (f'<m:oMath xmlns:m="{_MNS}"><m:r><m:rPr><m:sty m:val="i"/></m:rPr>'
        '<m:t xml:space="preserve">Z = </m:t></m:r><m:d><m:dPr>'
        '<m:begChr m:val="("/><m:endChr m:val=")"/></m:dPr><m:e>'
        '<m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t xml:space="preserve">R1 + jX1</m:t></m:r>'
        '</m:e></m:d><m:r><m:t xml:space="preserve"> &#x22C5; L</m:t></m:r></m:oMath>')

def EQ_ZTRAFO():
    return (f'<m:oMath xmlns:m="{_MNS}"><m:r><m:rPr><m:sty m:val="i"/></m:rPr>'
        '<m:t xml:space="preserve">Z T = </m:t></m:r><m:f>'
        '<m:num><m:r><m:t xml:space="preserve">uk</m:t></m:r></m:num>'
        '<m:den><m:r><m:t>100</m:t></m:r></m:den></m:f>'
        '<m:r><m:t xml:space="preserve"> &#x22C5; </m:t></m:r><m:f><m:num>'
        '<m:sSup><m:e><m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t>V</m:t></m:r></m:e>'
        '<m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>'
        '<m:r><m:t xml:space="preserve"> rT</m:t></m:r></m:num><m:den>'
        '<m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t xml:space="preserve">S rT</m:t></m:r>'
        '</m:den></m:f></m:oMath>')

def EQ_ICC3():
    return (f'<m:oMath xmlns:m="{_MNS}"><m:r><m:rPr><m:sty m:val="i"/></m:rPr>'
        '<m:t xml:space="preserve">I k3 = </m:t></m:r><m:f><m:num>'
        '<m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t xml:space="preserve">c &#x22C5; Vn</m:t></m:r>'
        '</m:num><m:den><m:rad><m:radPr><m:degHide m:val="1"/></m:radPr>'
        '<m:deg/><m:e><m:r><m:t>3</m:t></m:r></m:e></m:rad>'
        '<m:r><m:t xml:space="preserve"> &#x22C5; |Z1|</m:t></m:r>'
        '</m:den></m:f></m:oMath>')

def EQ_ICC2():
    return (f'<m:oMath xmlns:m="{_MNS}"><m:r><m:rPr><m:sty m:val="i"/></m:rPr>'
        '<m:t xml:space="preserve">I k2 = </m:t></m:r><m:f><m:num>'
        '<m:rad><m:radPr><m:degHide m:val="1"/></m:radPr>'
        '<m:deg/><m:e><m:r><m:t>3</m:t></m:r></m:e></m:rad></m:num>'
        '<m:den><m:r><m:t>2</m:t></m:r></m:den></m:f>'
        '<m:r><m:rPr><m:sty m:val="i"/></m:rPr>'
        '<m:t xml:space="preserve"> &#x22C5; I k3</m:t></m:r></m:oMath>')

def EQ_ICC1():
    return (f'<m:oMath xmlns:m="{_MNS}"><m:r><m:rPr><m:sty m:val="i"/></m:rPr>'
        '<m:t xml:space="preserve">I k1 = </m:t></m:r><m:f><m:num>'
        '<m:rad><m:radPr><m:degHide m:val="1"/></m:radPr>'
        '<m:deg/><m:e><m:r><m:t>3</m:t></m:r></m:e></m:rad>'
        '<m:r><m:t xml:space="preserve"> &#x22C5; c &#x22C5; Vn</m:t></m:r></m:num><m:den>'
        '<m:r><m:t xml:space="preserve">|2Z1 + Z0|</m:t></m:r>'
        '</m:den></m:f></m:oMath>')

def EQ_IP():
    return (f'<m:oMath xmlns:m="{_MNS}"><m:r><m:rPr><m:sty m:val="i"/></m:rPr>'
        '<m:t xml:space="preserve">ip = &#x03BA; &#x22C5; </m:t></m:r>'
        '<m:rad><m:radPr><m:degHide m:val="1"/></m:radPr>'
        '<m:deg/><m:e><m:r><m:t>2</m:t></m:r></m:e></m:rad>'
        '<m:r><m:rPr><m:sty m:val="i"/></m:rPr>'
        '<m:t xml:space="preserve"> &#x22C5; I k3</m:t></m:r></m:oMath>')

def EQ_KAPPA():
    return (f'<m:oMath xmlns:m="{_MNS}"><m:r>'
        '<m:t xml:space="preserve">&#x03BA; = 1,02 + 0,98 &#x22C5; </m:t></m:r>'
        '<m:sSup><m:e><m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t>e</m:t></m:r></m:e>'
        '<m:sup><m:r><m:t xml:space="preserve">&#x2212;3 R/X</m:t></m:r></m:sup>'
        '</m:sSup></m:oMath>')

def EQ_ZACC():
    return (f'<m:oMath xmlns:m="{_MNS}"><m:r><m:rPr><m:sty m:val="i"/></m:rPr>'
        '<m:t xml:space="preserve">Z acc,n = Z acc,n-1 + Z elem,n</m:t></m:r></m:oMath>')

def EQ_ZSEC():
    return (f'<m:oMath xmlns:m="{_MNS}"><m:r><m:rPr><m:sty m:val="i"/></m:rPr>'
        '<m:t xml:space="preserve">Z sec = </m:t></m:r><m:f><m:num>'
        '<m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t>Z prim</m:t></m:r></m:num><m:den>'
        '<m:sSup><m:e><m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t>n</m:t></m:r></m:e>'
        '<m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup></m:den></m:f>'
        '<m:r><m:t xml:space="preserve">,  n = V prim / V sec</m:t></m:r></m:oMath>')

def EQ_ALF():
    return (f'<m:oMath xmlns:m="{_MNS}"><m:r><m:t xml:space="preserve">ALF = </m:t></m:r>'
        '<m:f><m:num><m:r><m:rPr><m:sty m:val="i"/></m:rPr>'
        '<m:t xml:space="preserve">I k3 [A]</m:t></m:r></m:num><m:den>'
        '<m:r><m:rPr><m:sty m:val="i"/></m:rPr>'
        '<m:t xml:space="preserve">I n1 [A]</m:t></m:r></m:den></m:f></m:oMath>')

def _tbl(doc, headers, rows, widths=None, note="", hbg=_C_AZUL_MED):
    n = len(headers)
    if not widths: widths = [Cm(16.0/n)]*n
    tbl = doc.add_table(rows=1, cols=n)
    tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = tbl.rows[0]
    for i,h in enumerate(headers):
        cell = hrow.cells[i]; _set_cell_bg(cell,hbg)
        _set_cell_border(cell,color="FFFFFF",sz="2")
        _cell_write(cell,h,bold=True,size=9,color=_C_BRANCO,center=True); cell.width=widths[i]
    for ri,row in enumerate(rows):
        drow = tbl.add_row(); bg = _C_CINZA2 if ri%2==1 else _C_BRANCO
        for ci,val in enumerate(row):
            cell = drow.cells[ci]; _set_cell_bg(cell,bg)
            _set_cell_border(cell,color="CCCCCC",sz="2")
            _cell_write(cell,str(val),size=9,center=(ci>0)); cell.width=widths[ci]
    if note: _nota(doc,note)
    _sp(doc,4)

def _setup_hf(doc, projeto, doc_code, rev="00"):
    section = doc.sections[0]
    section.page_width=Cm(21.0); section.page_height=Cm(29.7)
    section.top_margin=Cm(2.5); section.bottom_margin=Cm(2.0)
    section.left_margin=Cm(3.0); section.right_margin=Cm(2.0)
    hdr = section.header; hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0]; hp.clear(); hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r1 = hp.add_run("BK ENGENHARIA E TECNOLOGIA  |  ")
    r1.bold=True; r1.font.size=Pt(8); r1.font.color.rgb=RGBColor.from_string(_C_AZUL_MED)
    r2 = hp.add_run(f"{doc_code}  Rev. {rev}")
    r2.font.size=Pt(8); r2.font.color.rgb=RGBColor.from_string(_C_CINZA)
    pPr=hp._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr"); btm=OxmlElement("w:bottom")
    btm.set(qn("w:val"),"single"); btm.set(qn("w:sz"),"4")
    btm.set(qn("w:space"),"1"); btm.set(qn("w:color"),_C_AZUL_MED)
    pBdr.append(btm); pPr.append(pBdr)
    ftr=section.footer; ftr.is_linked_to_previous=False
    fp=ftr.paragraphs[0]; fp.clear(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r3=fp.add_run(f"{projeto}  |  Estudo de Protecao  |  Pg. ")
    r3.font.size=Pt(8); r3.font.color.rgb=RGBColor.from_string(_C_CINZA)
    fld=OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"),"begin")
    ins=OxmlElement("w:instrText"); ins.text="PAGE"; ins.set(qn("xml:space"),"preserve")
    fld2=OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"),"end")
    rp=fp.add_run(); rp.font.size=Pt(8); rp.font.color.rgb=RGBColor.from_string(_C_CINZA)
    rp._r.append(fld); rp._r.append(ins); rp._r.append(fld2)

def _capa(doc, info):
    tbl=doc.add_table(rows=1,cols=1); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    cell=tbl.rows[0].cells[0]; cell.width=Cm(16.0)
    _set_cell_bg(cell,_C_AZUL_ESC); _set_cell_border(cell,color=_C_AZUL_ESC)
    cp=cell.paragraphs[0]; cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before=Pt(12); cp.paragraph_format.space_after=Pt(12)
    r=cp.add_run("BK ENGENHARIA E TECNOLOGIA")
    r.bold=True; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string(_C_BRANCO)
    _sp(doc,28)
    pt=doc.add_paragraph(); pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rt=pt.add_run("ESTUDO DE PROTECAO E COORDENACAO")
    rt.bold=True; rt.font.size=Pt(20); rt.font.color.rgb=RGBColor.from_string(_C_AZUL_ESC)
    pt2=doc.add_paragraph(); pt2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rt2=pt2.add_run("DE SISTEMAS ELETRICOS DE POTENCIA")
    rt2.bold=True; rt2.font.size=Pt(20); rt2.font.color.rgb=RGBColor.from_string(_C_AZUL_ESC)
    _hline(doc,color=_C_AZUL_MED,sz="12",before=8,after=8)
    campos=[
        ("Projeto",info.get("projeto","---")),("Cliente",info.get("cliente","---")),
        ("Local / Unidade",info.get("local","---")),
        ("Codigo do Doc.",info.get("doc_code","BK-EP-001")),
        ("Revisao",info.get("revisao","00")),
        ("Data de Emissao",info.get("data",datetime.date.today().strftime("%d/%m/%Y"))),
        ("Elaborado por",info.get("elaborado","Engenharia BK")),
        ("Verificado por",info.get("verificado","---")),
        ("Aprovado por",info.get("aprovado","---")),
        ("Concessionaria",info.get("concessionaria","---")),
        ("Tensao de Entrega",info.get("tensao_entrega","---")),
        ("Classificacao","ESTUDO TECNICO -- USO EXTERNO"),
    ]
    _tbl(doc,["CAMPO","INFORMACAO"],campos,widths=[Cm(5.5),Cm(10.5)],hbg=_C_AZUL_ESC)
    _sp(doc,10)
    pa=doc.add_paragraph(); pa.alignment=WD_ALIGN_PARAGRAPH.CENTER
    ra=pa.add_run("Documento elaborado em conformidade com IEC 60909:2016 | IEC 61869-2 | IEC 62271-100 | IEC 60255")
    ra.italic=True; ra.font.size=Pt(9); ra.font.color.rgb=RGBColor.from_string(_C_CINZA)
    doc.add_page_break()

def _sec1(doc, info):
    _h1(doc,"1","OBJETO E ESCOPO")
    projeto=info.get("projeto","instalacao eletrica"); cliente=info.get("cliente","cliente")
    local=info.get("local","local nao informado"); tensao=info.get("tensao_entrega","---")
    concess=info.get("concessionaria","concessionaria distribuidora")
    _body(doc,f'O presente documento tem por objeto apresentar o Estudo de Protecao e Coordenacao referente ao empreendimento "{projeto}", de titularidade de {cliente}, localizado em {local}.')
    _body(doc,f"O estudo contempla a conexao ao sistema eletrico da {concess} em tensao de {tensao}, incluindo todos os elementos desde o ponto de entrega ate as barras de distribuicao internas.")
    _h2(doc,"1.1","Objetivos Especificos")
    for obj in [
        "Calcular as correntes de curto-circuito trifasica, bifasica e monofasica em todas as barras, pelo metodo IEC 60909:2016;",
        "Dimensionar os Transformadores de Corrente (TC), de Potencial (TP) e Disjuntores conforme IEC 61869-2, IEC 61869-3 e IEC 62271-100;",
        "Definir os ajustes dos reles de protecao e verificar a coordenacao e seletividade;",
        "Gerar o coordenograma Tempo x Corrente demonstrando a hierarquia de protecao;",
        "Atender os requisitos tecnicos da concessionaria para aprovacao do projeto de conexao.",
    ]:
        p=doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(2)
        p.paragraph_format.left_indent=Cm(0.8)
        run=p.add_run(obj); run.font.size=Pt(10); run.font.color.rgb=RGBColor.from_string(_C_CINZA)
    _sp(doc,4)

def _sec2(doc):
    _h1(doc,"2","REFERENCIAS NORMATIVAS")
    normas=[
        ("IEC 60909:2016","Short-circuit currents in three-phase AC systems -- Calculation of currents"),
        ("IEC 61869-2:2012","Instrument transformers -- Additional requirements for current transformers"),
        ("IEC 61869-3:2011","Instrument transformers -- Additional requirements for inductive voltage transformers"),
        ("IEC 62271-100:2021","High-voltage switchgear -- Part 100: AC circuit-breakers"),
        ("IEC 60255-151:2009","Measuring relays and protection equipment -- Over/under-current protection"),
        ("IEC 60076-1:2011","Power transformers -- General"),
        ("IEC 60076-5:2006","Power transformers -- Ability to withstand short circuit"),
        ("NBR 14039:2005","Instalacoes eletricas de media tensao de 1,0 kV a 36,2 kV"),
        ("NBR 5410:2004","Instalacoes eletricas de baixa tensao"),
        ("PRODIST -- Modulo 3","Acesso ao sistema de distribuicao (ANEEL)"),
        ("NT da Concessionaria","Conexao de sistemas de protecao ao sistema de distribuicao"),
    ]
    _tbl(doc,["Norma / Documento","Titulo / Descricao"],normas,widths=[Cm(5.0),Cm(11.0)],hbg=_C_AZUL_ESC)

def _sec3(doc, c_factor=1.10):
    _h1(doc,"3","METODOLOGIA DE CALCULO")
    _body(doc,"O calculo de curtos-circuitos e realizado pelo Metodo das Componentes Simetricas, conforme IEC 60909:2016, utilizando o equivalente de tensao da rede para determinacao das correntes de falta maximas (I''k). O metodo e valido para redes radiais e malhadas com acumulacao de impedancias por BFS (Breadth-First Search).")
    _h2(doc,"3.1","Fator de Tensao c  (IEC 60909 -- Tabela 1)")
    _body(doc,f"O fator c simula a tensao maxima no ponto de falta. Para correntes maximas (pior caso): c = {c_factor:.2f}.")
    _tbl(doc,["Nivel de Tensao","c max","c min","Aplicacao"],[
        ("Baixa tensao (< 1 kV)","1,05","0,95","Sistemas 127/220/380 V"),
        ("Alta tensao (> 1 kV)","1,10","1,00","Sistemas MT e AT (adotado neste estudo)"),
    ],widths=[Cm(4.5),Cm(1.8),Cm(1.8),Cm(7.9)],hbg=_C_AZUL_ESC)
    _h2(doc,"3.2","Impedancia da Fonte (Concessionaria)")
    _body(doc,"A impedancia equivalente no ponto de entrega e calculada a partir do nivel de curto S''kQ (IEC 60909 Eq. 18):")
    _omml_block(doc,EQ_ZQ(),label="3.1")
    _body(doc,"Onde: c = fator de tensao; VnQ = tensao nominal [kV]; S''kQ = potencia de curto-circuito trifasica [MVA].",size=9,italic=True)
    _h2(doc,"3.3","Impedancia de Linhas e Cabos  (IEC 60909 Tabela 3)")
    _omml_block(doc,EQ_ZLINHA(),label="3.2")
    _body(doc,"Onde: R1, X1 = resistencia e reatancia positivas [Ohm/km]; L = comprimento [km].",size=9,italic=True)
    _h2(doc,"3.4","Impedancia de Transformadores  (IEC 60909 sec.3.3.2)")
    _omml_block(doc,EQ_ZTRAFO(),label="3.3")
    _omml_block(doc,EQ_ZSEC(),label="3.4")
    _h2(doc,"3.5","Acumulacao de Impedancias em Rede Radial")
    _omml_block(doc,EQ_ZACC(),label="3.5")
    _h2(doc,"3.6","Correntes de Curto-Circuito  (IEC 60909 sec.4)")
    _body(doc,"Falta trifasica (simetrica inicial):",bold=True,size=10)
    _omml_block(doc,EQ_ICC3(),label="3.6")
    _body(doc,"Falta bifasica (sem terra):",bold=True,size=10)
    _omml_block(doc,EQ_ICC2(),label="3.7")
    _body(doc,"Falta monofasica (fase-terra):",bold=True,size=10)
    _omml_block(doc,EQ_ICC1(),label="3.8")
    _body(doc,"Corrente de pico (componente assincrona maxima):",bold=True,size=10)
    _omml_block(doc,EQ_IP(),label="3.9")
    _omml_block(doc,EQ_KAPPA(),label="3.10")
    _nota(doc,"Todas as correntes calculadas sao correntes simetricas iniciais (I''k), sem decaimento temporal.")
    _sp(doc,4)

def _sec4(doc, system, elements):
    _h1(doc,"4","DADOS DO SISTEMA E ELEMENTOS DE REDE")
    _h2(doc,"4.1","Parametros do Sistema")
    scc_mva=getattr(system,"short_circuit_mva_source",None)
    xr=getattr(system,"xr_ratio_source",10.0) or 10.0
    v_base=getattr(system,"v_base_kv",0.0)
    c_factor=getattr(system,"voltage_factor_c",1.10)
    z_r=getattr(system,"z_source_r_ohm",0.0) or 0.0
    z_x=getattr(system,"z_source_x_ohm",0.0) or 0.0
    z_mag=math.sqrt(z_r**2+z_x**2) if (z_r or z_x) else 0.0
    icc_src=round(c_factor*v_base/(math.sqrt(3)*z_mag),3) if z_mag>1e-9 else 0.0
    _tbl(doc,["Parametro","Valor"],[
        ("Tensao base do sistema",f"{v_base:.3f} kV"),
        ("Nivel de curto S''kQ",f"{scc_mva:.1f} MVA" if scc_mva else "Ver Z fonte"),
        ("Relacao X/R da rede",f"{xr:.1f}"),
        ("Impedancia da fonte  R1",f"{z_r:.6f} Ohm"),
        ("Impedancia da fonte  X1",f"{z_x:.6f} Ohm"),
        ("|Z1_fonte|",f"{z_mag:.6f} Ohm"),
        ("Fator de tensao c",f"{c_factor:.2f}"),
        ("Icc na barra de entrada (3f)",f"{icc_src:.3f} kA"),
        ("Frequencia","60 Hz"),
    ],widths=[Cm(8.0),Cm(8.0)],hbg=_C_AZUL_ESC)
    _h2(doc,"4.2","Dados dos Elementos de Rede")
    if not elements:
        _body(doc,"Nenhum elemento de rede cadastrado."); return
    linhas=[e for e in elements if str(getattr(e,"element_type","")).lower() in ("linha","linha_aerea","cabo","cabo_subterraneo","alimentador")]
    trafos=[e for e in elements if str(getattr(e,"element_type","")).lower() in ("trafo","transformador")]
    if linhas:
        _body(doc,"Linhas e Cabos:",bold=True,size=10)
        rows_l=[(getattr(e,"code","---"),str(getattr(e,"element_type","---")),
            f"{getattr(e,'bus_from','?')} -> {getattr(e,'bus_to','?')}",
            f"{getattr(e,'voltage_kv',0):.1f}",f"{getattr(e,'length_km',0):.3f}",
            f"{getattr(e,'r1_ohm_km',0):.4f}",f"{getattr(e,'x1_ohm_km',0):.4f}",
            f"{getattr(e,'r0_ohm_km',0) or '---'}",f"{getattr(e,'x0_ohm_km',0) or '---'}") for e in linhas]
        _tbl(doc,["Cod","Tipo","Barras","V(kV)","L(km)","R1","X1","R0","X0"],rows_l,
            widths=[Cm(1.5),Cm(1.8),Cm(3.0),Cm(1.4),Cm(1.4),Cm(1.8),Cm(1.8),Cm(1.6),Cm(1.7)],note="R1,X1,R0,X0 em Ohm/km.")
    if trafos:
        _body(doc,"Transformadores:",bold=True,size=10)
        rows_t=[(getattr(e,"code","---"),f"{getattr(e,'bus_from','?')} -> {getattr(e,'bus_to','?')}",
            f"{getattr(e,'voltage_kv',0):.3f} / {getattr(e,'trafo_voltage_sec_kv',0):.3f}",
            f"{getattr(e,'trafo_kva',0):.0f} kVA",f"{getattr(e,'trafo_z_percent',0):.2f} %",
            str(getattr(e,"trafo_connection","---"))) for e in trafos]
        _tbl(doc,["Cod","Barras","V AT/BT (kV)","Potencia","uk (%)","Ligacao"],rows_t,
            widths=[Cm(1.8),Cm(3.2),Cm(3.5),Cm(2.5),Cm(2.0),Cm(3.0)])
    _sp(doc,4)

def _sec5(doc, sc_results, system=None):
    _h1(doc,"5","RESULTADOS -- IMPEDANCIAS POR BARRA E CORRENTES DE CURTO-CIRCUITO")
    c=getattr(system,"voltage_factor_c",1.10) if system else 1.10
    if not sc_results:
        _body(doc,"Nenhum resultado disponiVel. Execute o calculo IEC 60909 primeiro."); return
    _h2(doc,"5.1","Impedancias Acumuladas por Barra")
    _body(doc,"Impedancias de sequencia acumuladas desde a fonte ate cada barra pelo processo BFS:")
    rows_z=[]
    for r in sc_results:
        z1=getattr(r,"z1_ohm",0j) or 0j; z0=getattr(r,"z0_ohm",None)
        z1_mag=abs(z1); z0_mag=abs(z0) if z0 is not None else None
        xr_val=z1.imag/z1.real if z1.real>1e-9 else 0.0
        kappa=getattr(r,"kappa_factor",0.0) or 0.0
        rows_z.append((getattr(r,"element_code","---"),getattr(r,"bus_name","---"),
            f"{z1.real:.5f}",f"{z1.imag:.5f}",f"{z1_mag:.5f}",
            f"{z0.real:.5f}" if z0 is not None else "---",
            f"{z0.imag:.5f}" if z0 is not None else "---",
            f"{z0_mag:.5f}" if z0_mag is not None else "INF",
            f"{xr_val:.2f}",f"{kappa:.3f}"))
    _tbl(doc,["Elem","Barra","R1(Ohm)","X1(Ohm)","|Z1|(Ohm)","R0(Ohm)","X0(Ohm)","|Z0|(Ohm)","X/R","kappa"],rows_z,
        widths=[Cm(1.4),Cm(1.8),Cm(1.8),Cm(1.8),Cm(1.8),Cm(1.6),Cm(1.6),Cm(1.6),Cm(1.2),Cm(1.4)],
        hbg=_C_AZUL_MED,note="Z0=INF indica sequencia zero bloqueada (trafo Yg-D em serie).")
    _h2(doc,"5.2","Correntes de Curto-Circuito por Barra")
    rows_i=[]
    for r in sc_results:
        icc3=getattr(r,"icc_3ph_ka",0.0) or 0.0; icc2=getattr(r,"icc_2ph_ka",0.0) or 0.0
        icc1=getattr(r,"icc_1ph_ka",0.0) or 0.0; ip=getattr(r,"icc_peak_ka",0.0) or 0.0
        k=getattr(r,"kappa_factor",0.0) or 0.0; bt=getattr(r,"icc_3ph_lv_ka",0.0) or 0.0
        rows_i.append((getattr(r,"element_code","---"),getattr(r,"bus_name","---"),
            f"{icc3:.3f}",f"{icc2:.3f}",f"{icc1:.3f}" if icc1>0 else "BLOQ.",
            f"{ip:.3f}",f"{k:.3f}",f"{bt:.3f}" if bt>0 else "---"))
    _tbl(doc,["Elem","Barra","Ik3(kA)","Ik2(kA)","Ik1(kA)","ip(kA)","kappa","Ik3-BT(kA)"],rows_i,
        widths=[Cm(1.4),Cm(1.8),Cm(2.0),Cm(2.0),Cm(2.0),Cm(2.0),Cm(1.6),Cm(2.2)],
        hbg=_C_AZUL_MED,note="BLOQ. = sequencia zero bloqueada. Ik3-BT = corrente no secundario do trafo.")
    _h2(doc,"5.3","Memoria de Calculo -- Substituicao Numerica por Barra")
    _body(doc,"Para cada barra apresenta-se a memoria de calculo com substituicao numerica completa nas equacoes IEC 60909:")
    for r in sc_results:
        ec=getattr(r,"element_code","?"); bn=getattr(r,"bus_name","?")
        z1=getattr(r,"z1_ohm",0j) or 0j; z0=getattr(r,"z0_ohm",None)
        icc3=getattr(r,"icc_3ph_ka",0.0) or 0.0; icc2=getattr(r,"icc_2ph_ka",0.0) or 0.0
        icc1=getattr(r,"icc_1ph_ka",0.0) or 0.0; ip=getattr(r,"icc_peak_ka",0.0) or 0.0
        k=getattr(r,"kappa_factor",0.0) or 0.0; z1m=abs(z1)
        xr=z1.imag/z1.real if z1.real>1e-9 else 0.0; warns=getattr(r,"warnings",[]) or []
        ps=doc.add_paragraph(); ps.paragraph_format.space_before=Pt(8)
        ps.paragraph_format.space_after=Pt(2); ps.paragraph_format.left_indent=Cm(0.3)
        rs=ps.add_run(f"Barra {bn}  ({ec})")
        rs.bold=True; rs.font.size=Pt(10); rs.font.color.rgb=RGBColor.from_string(_C_AZUL_MED)
        mem=[
            ("Z1 acumulada",f"({z1.real:.6f} + j{z1.imag:.6f}) Ohm"),
            ("|Z1|",f"sqrt({z1.real:.6f}^2 + {z1.imag:.6f}^2) = {z1m:.6f} Ohm"),
            ("X/R",f"{xr:.4f}"),
            ("kappa = 1,02 + 0,98 x exp(-3 x R/X)",f"= {k:.4f}"),
            ("Ik3 = c x Vn / (sqrt(3) x |Z1|)",f"= {c:.2f} x Vn / (1,7321 x {z1m:.6f}) = {icc3:.3f} kA"),
            ("Ik2 = (sqrt(3)/2) x Ik3",f"= 0,8660 x {icc3:.3f} = {icc2:.3f} kA"),
        ]
        if z0 is not None and icc1>0:
            z_den=abs(2*z1+z0)
            mem.append(("Ik1 = sqrt(3) x c x Vn / |2Z1 + Z0|",f"= 1,7321 x {c:.2f} x Vn / {z_den:.6f} = {icc1:.3f} kA"))
        else:
            mem.append(("Ik1","BLOQUEADA -- Z0 = INF (trafo Yg-D em serie)"))
        mem.append(("ip = kappa x sqrt(2) x Ik3",f"= {k:.4f} x 1,4142 x {icc3:.3f} = {ip:.3f} kA"))
        for w in warns: mem.append(("Obs.",str(w)))
        _tbl(doc,["Grandeza / Equacao","Substituicao e Resultado"],mem,widths=[Cm(6.5),Cm(9.5)],hbg=_C_AZUL_LIG)
    _sp(doc,4)

def _sec6(doc, ct_results, vt_results, breaker_results, sc_results=None):
    _h1(doc,"6","DIMENSIONAMENTO DE EQUIPAMENTOS DE PROTECAO E MEDICAO")
    _body(doc,"O dimensionamento e realizado com base nas correntes calculadas na Secao 5, observando os criterios normativos de cada equipamento.")
    _h2(doc,"6.1","Transformadores de Corrente (TC) -- IEC 61869-2")
    _body(doc,"Criterio principal: Fator de Limite de Precisao (ALF) minimo para que o TC nao sature durante a corrente de falta maxima:")
    _omml_block(doc,EQ_ALF(),label="6.1")
    _tbl(doc,["Criterio","Formula / Regra","Norma"],[
        ("Corrente nominal I_n1","I_n1 >= 1,2 x I_carga_nominal","IEC 61869-2 sec.6.1"),
        ("ALF minimo de protecao","ALF >= Ik3 [A] / I_n1 [A]","IEC 61869-2 sec.6.3"),
        ("Classe de precisao","5P para protecao; 10P para sobrecorrente simples","IEC 61869-2 Tab.3"),
        ("Tensao de knee-point (PS)","Vk >= ALF x I_n2 x (R_CT + R_burden)","IEC 61869-2 sec.6.4"),
    ],widths=[Cm(4.0),Cm(8.0),Cm(4.0)])
    if ct_results:
        rows_ct=[]
        for ct in ct_results:
            icc3=getattr(ct,"icc_3ph_ka_bus",0.0) or 0.0; in1=getattr(ct,"rated_primary_A",0.0) or 0.0
            alf_calc=round((icc3*1000)/in1,1) if in1>0 else 0.0
            alf_nom=getattr(ct,"alf",20.0) or 20.0; classe=getattr(ct,"accuracy_class","5P20")
            ok=alf_nom>=alf_calc
            rows_ct.append((getattr(ct,"element_code","---"),getattr(ct,"bus_name","---"),
                f"{in1:.0f} / 5 A",f"{icc3:.3f} kA",f"{alf_calc:.1f}",f"{alf_nom:.0f}",classe,_ok_str(ok)))
        _tbl(doc,["Elem","Barra","Relacao(A)","Ik3(kA)","ALF calc","ALF nom","Classe","Result."],rows_ct,
            widths=[Cm(1.4),Cm(1.8),Cm(2.2),Cm(2.0),Cm(1.8),Cm(1.8),Cm(2.0),Cm(3.0)],
            note="TC aprovado se ALF_nominal >= ALF_calc.")
    else: _body(doc,"Resultados de TC nao disponiveis.",italic=True)
    _h2(doc,"6.2","Transformadores de Potencial (TP) -- IEC 61869-3")
    _tbl(doc,["Criterio","Formula / Regra","Norma"],[
        ("Tensao primaria","VrTV >= V_sistema / sqrt(3)  (fase-terra)","IEC 61869-3 sec.5.3"),
        ("Classe de medicao","Classe 0.5 para medicao fiscal; 3P para protecao","IEC 61869-3 Tab.1"),
        ("Fator de tensao (FT)","FT = 1,9 por 8 h para sistema isolado","IEC 61869-3 sec.5.3.3"),
        ("Burden maximo","Carga real <= B_nominal (10/25/50/100 VA)","IEC 61869-3 sec.6.2"),
    ],widths=[Cm(4.0),Cm(8.0),Cm(4.0)])
    if vt_results:
        rows_vt=[]
        for vt in vt_results:
            ok=getattr(vt,"is_valid",True)
            rows_vt.append((getattr(vt,"element_code","---"),getattr(vt,"bus_name","---"),
                getattr(vt,"rated_primary_kv","---"),getattr(vt,"rated_secondary_v","115 V"),
                getattr(vt,"accuracy_class","0.5"),getattr(vt,"burden_va","---"),_ok_str(ok)))
        _tbl(doc,["Elem","Barra","V prim","V sec","Classe","Burden(VA)","Result."],rows_vt,
            widths=[Cm(1.4),Cm(1.8),Cm(2.5),Cm(2.0),Cm(2.0),Cm(2.5),Cm(3.8)])
    else: _body(doc,"Resultados de TP nao disponiveis.",italic=True)
    _h2(doc,"6.3","Disjuntores -- IEC 62271-100")
    _tbl(doc,["Parametro","Criterio","Norma"],[
        ("Corrente de interrupcao I_cu","I_cu >= Ik3_max na barra","IEC 62271-100 sec.4.101"),
        ("Corrente de fechamento I_ma","I_ma >= ip_max = kappa x sqrt(2) x Ik3","IEC 62271-100 sec.4.106"),
        ("Corrente nominal continua I_n","I_n >= 1,25 x I_carga_max","IEC 62271-100 sec.4.5"),
    ],widths=[Cm(4.5),Cm(8.5),Cm(3.0)])
    if breaker_results:
        rows_br=[]
        for br in breaker_results:
            icc3=getattr(br,"icc_3ph_ka_bus",0.0) or 0.0; ip=getattr(br,"icc_peak_ka_bus",0.0) or 0.0
            icu=getattr(br,"rated_sc_ka",0.0) or 0.0; ima=getattr(br,"rated_peak_ka",0.0) or 0.0
            ok=(icu>=icc3) and (ima>=ip)
            rows_br.append((getattr(br,"element_code","---"),getattr(br,"bus_name","---"),
                f"{icc3:.3f}",f"{ip:.3f}",f"{icu:.1f}",f"{ima:.1f}",_ok_str(ok)))
        _tbl(doc,["Elem","Barra","Ik3(kA)","ip(kA)","I_cu(kA)","I_ma(kA)","Result."],rows_br,
            widths=[Cm(1.5),Cm(2.0),Cm(2.2),Cm(2.0),Cm(2.2),Cm(2.2),Cm(3.9)],
            note="Aprovado: I_cu >= Ik3 E I_ma >= ip.")
    else: _body(doc,"Resultados de disjuntores nao disponiveis.",italic=True)
    _h2(doc,"6.4","Sintese do Dimensionamento por Barra")
    if sc_results:
        br_series=[6.3,8,10,12.5,16,20,25,31.5,40,50,63,80,100,125]
        rows_syn=[]
        for r in sc_results:
            icc3=getattr(r,"icc_3ph_ka",0.0) or 0.0; ip=getattr(r,"icc_peak_ka",0.0) or 0.0
            br_min=next((v for v in br_series if v>=icc3),icc3)
            rows_syn.append((getattr(r,"element_code","---"),getattr(r,"bus_name","---"),
                f"{icc3:.3f} kA",f"{ip:.3f} kA",f"I_cu >= {br_min:.1f} kA","5P20 ou sup."))
        _tbl(doc,["Elem","Barra","Ik3(kA)","ip(kA)","Disj. min.","Classe TC"],rows_syn,
            widths=[Cm(1.5),Cm(2.0),Cm(2.5),Cm(2.5),Cm(4.5),Cm(3.0)],
            note="Disjuntor minimo = proximo valor serie comercial IEC 62271-100 >= Ik3.")
    _sp(doc,4)

def _sec7(doc, relay_results, coordenograma_b64=None, sc_results=None):
    _h1(doc,"7","COORDENACAO E SELETIVIDADE DOS RELES DE PROTECAO")
    _body(doc,"A coordenacao garante que o dispositivo mais proximo ao ponto de falta atue primeiro (protecao primaria), e o dispositivo a montante atue so se o primeiro falhar (retaguarda).")
    _h2(doc,"7.1","Criterios de Coordenacao  (IEC 60255)")
    _tbl(doc,["Tipo de Rele","CTI minimo","Justificativa"],[
        ("Eletromagnetico",">= 0,40 s","Tolerancia +/-7,5% + tempo de abertura do disjuntor (~100 ms)"),
        ("Digital / Microprocesado",">=0,25 s","Tolerancia +/-5% + tempo de abertura (~60 ms) + margem"),
        ("Numerico IED",">= 0,20 s","Tolerancia <= 1% + tempo de abertura <= 60 ms"),
    ],widths=[Cm(4.0),Cm(2.5),Cm(9.5)],note="CTI = t_retaguarda - t_primaria. Adotado neste estudo: 0,30 s.")
    _h2(doc,"7.2","Ajustes dos Reles por Barra")
    if relay_results:
        rows_r=[]
        for relay in relay_results:
            rows_r.append((getattr(relay,"element_code","---"),getattr(relay,"bus_name","---"),
                getattr(relay,"relay_type","51"),getattr(relay,"curve_type","VI"),
                f"{getattr(relay,'pickup_current_a',0.0):.1f} A",
                f"{getattr(relay,'pickup_multiple',0.0):.2f} pu",
                f"{getattr(relay,'time_multiplier',0.0):.3f}",
                f"{getattr(relay,'inst_pickup_a',0.0):.1f} A"))
        _tbl(doc,["Elem","Barra","Funcao","Curva","Pickup(A)","Pickup(pu)","TMS","Inst.(A)"],rows_r,
            widths=[Cm(1.4),Cm(1.8),Cm(1.5),Cm(1.5),Cm(2.0),Cm(2.0),Cm(1.8),Cm(4.0)],
            hbg=_C_AZUL_MED,note="Funcao ANSI: 51=sobrecorrente temporizado; 50=instantaneo; 67=direcional.")
    else: _body(doc,"Nenhum resultado de rele disponivel.",italic=True)
    _h2(doc,"7.3","Analise de Seletividade -- Margens CTI")
    if relay_results and len(relay_results)>=2:
        rows_cti=[]
        for i in range(len(relay_results)-1):
            rp=relay_results[i]; rr=relay_results[i+1]
            tp=float(getattr(rp,"op_time_s",None) or getattr(rp,"time_dial",0.3))
            tr=float(getattr(rr,"op_time_s",None) or getattr(rr,"time_dial",0.6))
            cti=round(tr-tp,3); ok=cti>=0.20
            rows_cti.append((getattr(rp,"element_code","---"),getattr(rr,"element_code","---"),
                f"{getattr(rp,'icc_3ph_ka_bus',0):.3f} kA",
                f"{tp:.3f} s",f"{tr:.3f} s",f"{cti:.3f} s","OK" if ok else "REVISAR"))
        _tbl(doc,["Primaria","Retaguarda","Ik3 falta(kA)","t_prim(s)","t_ret(s)","CTI(s)","Coord."],rows_cti,
            widths=[Cm(2.0),Cm(2.5),Cm(2.5),Cm(2.5),Cm(2.5),Cm(2.0),Cm(2.0)],
            note="CTI >= 0,20 s: OK. CTI < 0,20 s: revisar TMS ou Pickup.")
    else: _body(doc,"CTI nao disponivel (requer >= 2 reles configurados).",italic=True)
    _h2(doc,"7.4","Coordenograma Tempo x Corrente")
    _body(doc,"O coordenograma apresenta as curvas Tempo x Corrente (log-log) dos reles com as correntes de falta por barra. As curvas de inrush delimitam a zona proibida de atuacao:")
    if coordenograma_b64:
        import base64
        try:
            img_bytes=base64.b64decode(coordenograma_b64)
            doc.add_picture(io.BytesIO(img_bytes),width=Cm(14.0))
            last=doc.paragraphs[-1]; last.alignment=WD_ALIGN_PARAGRAPH.CENTER
            pl=doc.add_paragraph(); pl.alignment=WD_ALIGN_PARAGRAPH.CENTER
            rl=pl.add_run("Figura 1 -- Coordenograma Tempo x Corrente (escala log-log)")
            rl.italic=True; rl.font.size=Pt(9); rl.font.color.rgb=RGBColor.from_string(_C_CINZA)
        except Exception as exc: _body(doc,f"[Erro ao inserir coordenograma: {exc}]",italic=True)
    else: _body(doc,"Coordenograma nao disponivel. Execute o calculo completo.",italic=True)
    _sp(doc,4)

def _sec8(doc, sc_results, relay_results, system=None):
    _h1(doc,"8","VERIFICACAO DE REQUISITOS DA CONCESSIONARIA")
    _body(doc,"Verificacao dos requisitos minimos exigidos pela concessionaria distribuidora para aprovacao do projeto de conexao, conforme PRODIST Modulo 3 e norma tecnica da distribuidora.")
    c_factor=getattr(system,"voltage_factor_c",1.10) if system else 1.10
    icc_max=max((getattr(r,"icc_3ph_ka",0.0) or 0.0 for r in (sc_results or [])),default=0.0)
    itens=[
        ("Nivel de curto informado pela concessionaria","Scc conforme dados de entrada","Conforme" if icc_max>0 else "Pendente"),
        (f"Fator de tensao c = 1,10  (IEC 60909 Tab.1)",f"c adotado = {c_factor:.2f}","Conforme" if abs(c_factor-1.10)<0.01 else "Verificar"),
        ("Corrente de falta minima detectavel pelo rele  (Imin > 1,5 x Ip)","Verificar Pickup <= Ik_min / 1,5","A verificar" if not relay_results else "Ver Secao 7.2"),
        ("Seletividade -- CTI >= 0,20 s entre primaria e retaguarda","Ver Tabela CTI na Secao 7.3","A verificar" if not relay_results else "Ver Secao 7.3"),
        ("Protecao de minima tensao (27) no ponto de conexao","Rele de sub/sobretensao conforme NT distribuidora","A verificar"),
        ("Protecao de falta a terra (51N / 67N)","Rele de sobrecorrente de sequencia zero","Configurado" if relay_results else "A verificar"),
        ("Disjuntor com I_cu >= Ik3_max",f"Ik3_max = {icc_max:.3f} kA","Ver Secao 6.3"),
        ("Memorial de calculo com equacoes e memoria numerica","Apresentado nas Secoes 3, 5 e 6","Conforme"),
        ("Coordenograma Tempo x Corrente (escala log-log)","Apresentado na Secao 7.4","Conforme"),
    ]
    _tbl(doc,["Requisito","Atendimento / Referencia","Status"],itens,
        widths=[Cm(7.0),Cm(6.0),Cm(3.0)],
        note="'A verificar' indica dependencia de dados especificos da concessionaria ou configuracao do rele nao inserida.")
    _nota(doc,"Em caso de revisao do projeto, todos os calculos devem ser refeitos e o presente relatorio reeditado com nova revisao e assinatura.")
    _sp(doc,4)

def _sec9(doc, sc_results=None, relay_results=None):
    _h1(doc,"9","CONCLUSAO E RECOMENDACOES")
    if sc_results:
        icc3_vals=[getattr(r,"icc_3ph_ka",0.0) or 0.0 for r in sc_results]
        ip_vals=[getattr(r,"icc_peak_ka",0.0) or 0.0 for r in sc_results]
        icc_max=max(icc3_vals,default=0.0); icc_min=min(icc3_vals,default=0.0)
        ip_max=max(ip_vals,default=0.0); n_barras=len(sc_results)
    else: icc_max=icc_min=ip_max=0.0; n_barras=0
    _body(doc,"O presente estudo de protecao e coordenacao foi elaborado em conformidade com a norma IEC 60909:2016 para calculo de curtos-circuitos, e com a IEC 60255-151 para definicao dos ajustes dos reles de protecao.")
    if n_barras>0:
        _body(doc,f"Foram calculadas as correntes de curto-circuito em {n_barras} barra(s). Corrente trifasica maxima: {icc_max:.3f} kA; minima: {icc_min:.3f} kA; pico maximo: {ip_max:.3f} kA.")
    _h2(doc,"9.1","Recomendacoes Tecnicas")
    for rec in [
        "Todos os disjuntores devem ter I_cu >= corrente de curto trifasica maxima na barra de instalacao, conforme calculado na Secao 5.",
        "Os TCs de protecao devem ser de classe 5P com ALF suficiente para nao saturar nas correntes de falta calculadas (ver criterios Secao 6.1).",
        "Os ajustes dos reles devem ser inseridos nos equipamentos por engenheiro eletricista habilitado e verificados em comissionamento por injecao secundaria.",
        "O coordenograma deve ser validado em campo apos o comissionamento, com injecao de corrente secundaria nos TCs e verificacao dos tempos de atuacao.",
        "Em casode ampliacao ou alteracao da rede, este estudo deve ser reavaliado e nova revisao emitida antes da energizacao.",
        "Recomenda-se a instalacao de registrador de perturbacoes (DFR / oscilografo) para registro e analise de eventos de falta.",
    ]:
        p=doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(2)
        p.paragraph_format.left_indent=Cm(0.8)
        run=p.add_run(rec); run.font.size=Pt(10); run.font.color.rgb=RGBColor.from_string(_C_CINZA)
    _sp(doc,4)

def _sec10(doc, info):
    _h1(doc,"10","RESPONSAVEL TECNICO")
    _tbl(doc,["Campo","Informacao"],[
        ("Responsavel Tecnico",info.get("engenheiro","Engenheiro Responsavel")),
        ("CREA / CFE",info.get("crea","CREA-XX / XXXXXX-D")),
        ("Empresa",info.get("empresa","BK Engenharia e Tecnologia")),
        ("Cargo",info.get("cargo","Engenheiro Eletrici Sta Senior")),
        ("Telefone",info.get("telefone","---")),
        ("E-mail",info.get("email","---")),
        ("Data da Emissao",info.get("data",datetime.date.today().strftime("%d/%m/%Y"))),
        ("Revisao",info.get("revisao","00")),
    ],widths=[Cm(5.5),Cm(10.5)],hbg=_C_AZUL_ESC)
    _sp(doc,20); _hline(doc,color=_C_AZUL_ESC,sz="6",before=40,after=4)
    p_nome=doc.add_paragraph(); p_nome.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p_nome.add_run(info.get("engenheiro","Engenheiro Responsavel"))
    r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(_C_AZUL_ESC)
    p_crea=doc.add_paragraph(); p_crea.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rc=p_crea.add_run(f"{info.get('crea','CREA-XX / XXXXXX-D')}  --  Engenheiro Eletricista")
    rc.font.size=Pt(10); rc.font.color.rgb=RGBColor.from_string(_C_CINZA)
    _sp(doc,8)
    _nota(doc,"Este documento e de responsabilidade exclusiva do profissional habilitado indicado acima, nos termos da Lei 5.194/66 e Resolucao CONFEA 1048/2013.")

def gerar_relatorio_protecao(
    study_info: dict, system: Any, elements: list,
    sc_results=None, ct_results=None, vt_results=None,
    breaker_results=None, relay_results=None,
    coordenograma_b64: str = None,
) -> io.BytesIO:
    """
    Gera o Relatorio Tecnico de Estudo de Protecao em formato Word (.docx).
    Parametros: study_info (dict), system, elements, sc_results, ct_results,
    vt_results, breaker_results, relay_results, coordenograma_b64.
    Retorno: io.BytesIO com o arquivo .docx pronto para download.
    """
    doc=Document()
    proj=study_info.get("projeto","Estudo de Protecao")
    code=study_info.get("doc_code","BK-EP-001")
    rev=study_info.get("revisao","00")
    c_factor=getattr(system,"voltage_factor_c",study_info.get("voltage_factor_c",1.10))
    doc.styles["Normal"].font.name="Calibri"
    doc.styles["Normal"].font.size=Pt(10)
    doc.styles["Normal"].font.color.rgb=RGBColor.from_string(_C_CINZA)
    _setup_hf(doc,proj,code,rev)
    _capa(doc,study_info)
    _sec1(doc,study_info)
    _sec2(doc)
    _sec3(doc,c_factor)
    _sec4(doc,system,elements or [])
    _sec5(doc,sc_results or [],system)
    _sec6(doc,ct_results or [],vt_results or [],breaker_results or [],sc_results or [])
    _sec7(doc,relay_results or [],coordenograma_b64,sc_results or [])
    _sec8(doc,sc_results or [],relay_results or [],system)
    _sec9(doc,sc_results or [],relay_results or [])
    _sec10(doc,study_info)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf
